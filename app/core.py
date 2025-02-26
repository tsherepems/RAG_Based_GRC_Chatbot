"""
app/core.py
Handles file operations and basic document processing (load, chunk).
"""

import os
import logging
from pathlib import Path
import json
import hashlib

from docx import Document as DocxDocument
import pymupdf4llm
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

DATA_DIR = "./Data/source"

# Path where the processed file cache will be stored.
PROCESSED_CACHE_PATH = "./Data/processed_files.json"

def load_processed_cache():
    """Load the processed files cache from a JSON file."""
    if os.path.exists(PROCESSED_CACHE_PATH):
        with open(PROCESSED_CACHE_PATH, "r") as f:
            return json.load(f)
    return {}

def save_processed_cache(cache):
    """Save the processed files cache to a JSON file."""
    with open(PROCESSED_CACHE_PATH, "w") as f:
        json.dump(cache, f)

def get_file_hash(file_path: str) -> str:
    """Compute a hash (MD5) for the given file."""
    hasher = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hasher.update(chunk)
    return hasher.hexdigest()

#save_uploaded_file
def save_uploaded_file(file, session_dir: str) -> str:
    """
    Saves an uploaded file into a session-specific directory.
    """
    os.makedirs(session_dir, exist_ok=True)
    file_path = Path(session_dir) / file.name
    with open(file_path, "wb") as f:
        f.write(file.read())
    logger.info(f"File saved: {file_path}")
    return str(file_path)

def load_file(file_path: str) -> str:
    """
    Loads content from a file based on its extension (.txt, .docx, .pdf).
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == ".docx":
        try:
            doc = DocxDocument(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.error(f"Error reading .docx file {file_path}: {e}")
            return ""

    elif ext == ".pdf":
        try:
            return pymupdf4llm.to_markdown(str(file_path))
        except Exception as e:
            logger.error(f"Error reading .pdf file {file_path} with pymupdf4llm: {e}")
            return ""

    else:
        raise ValueError(f"Unsupported file format: {ext}")

def load_existing_files(source_dir: str = DATA_DIR):
    """
    Loads all existing files in the source directory as LangChain Documents.
    """
    documents = []
    for path in Path(source_dir).rglob("*"):
        if path.is_file():
            try:
                text = load_file(str(path))
                if text.strip():
                    documents.append(Document(page_content=text, metadata={"source": str(path)}))
            except ValueError as ve:
                logger.warning(f"Skipping file {path}: {ve}")
    logger.info(f"Loaded {len(documents)} documents from {source_dir}")
    return documents



def chunk_documents(documents, chunk_size=1000, chunk_overlap=200):
    """
    Splits documents into manageable chunks.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_documents(documents)
    logger.info(f"Chunked {len(documents)} documents into {len(chunks)} chunks.")
    return chunks
